#!/usr/bin/env python3
"""
generate_pack.py — build a District Consumer Commission filing pack from a case YAML.

Produces six .docx files (Index, Proforma, Synopsis+Dates, Memo of Parties,
Consumer Complaint+Affidavit, Party-in-Person Declaration) in the standard
Indian consumer-court house style.

The generator owns STRUCTURE: headings, paragraph numbering, the six-file split,
Verification/Affidavit boilerplate, Index rows, and every cross-reference
("paragraphs 1 to N", "Annexures A-1 to A-N"). The case YAML owns SUBSTANCE.

Usage:
    py -3.11 generate_pack.py <case.yaml> [--out <dir>]

Requires: python-docx, pyyaml
"""
import sys
import os
import argparse

try:
    import yaml
except ImportError:
    sys.exit("Missing dependency: pyyaml.  Install with:  py -3.11 -m pip install pyyaml")
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Missing dependency: python-docx.  Install with:  py -3.11 -m pip install python-docx")

# Conventional Index page ranges. Real pagination varies; these match the
# observed house convention and should be adjusted after the final print.
INDEX_PAGE_RANGES = {
    "index": "1",
    "proforma": "2 – 4",
    "synopsis": "5 – 7",
    "memo": "8",
    "complaint": "9 – 18",
    "affidavit": "19 – 20",
}
ROMAN = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x", "xi", "xii"]


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------
def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    return doc


def para(doc, text="", bold=False, italic=False, align=None, size=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    return p


def cause_title(doc, case):
    """The two centred bold lines at the head of every document."""
    c = case["commission"]
    line1 = "BEFORE THE {}, {}".format(c["name"], c["place"])
    para(doc, line1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "CONSUMER COMPLAINT NO. ______ OF {}".format(case["filing"]["year"]),
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def signature_block(doc, case, role="Complainant"):
    """Place / Date (day blank) / signature."""
    month = case["filing"].get("signing_month", "") or ""
    year = case["filing"]["year"]
    date_tail = ("{} {}".format(month, year)) if month else str(year)
    para(doc, "Place: {}".format(case["complainant"]["city"]), space_after=2)
    para(doc, "Date: ___________ {}".format(date_tail), space_after=12)
    para(doc, "({})".format(case["complainant"]["name"]),
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(doc, role, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)


def op_label(idx, total):
    return "… OPPOSITE PARTY" if total == 1 else "… OPPOSITE PARTY NO. {}".format(idx)


def party_block(doc, case):
    """IN THE MATTER OF ... COMPLAINANT / VERSUS / OPPOSITE PARTIES."""
    cm = case["complainant"]
    para(doc, "IN THE MATTER OF:", bold=True, space_after=8)
    para(doc, "{}, aged {} years,".format(cm["name"], cm["age"]), space_after=0)
    para(doc, "S/o Mr. {}".format(cm["father_name"]), space_after=0)
    para(doc, "R/o {}".format(cm["address"]), space_after=0)
    para(doc, "Email: {} | Mobile: {}".format(cm["email"], cm["mobile"]), space_after=0)
    para(doc, "… COMPLAINANT", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    para(doc, "VERSUS", bold=True, space_after=6)

    ops = case["opposite_parties"]
    for i, op in enumerate(ops, start=1):
        head = "{}. {}".format(i, op["name"])
        if op.get("operating_as"):
            head += " ({})".format(op["operating_as"])
        para(doc, head, space_after=0)
        if not op.get("is_individual") and op.get("cin"):
            para(doc, "CIN: {}".format(op["cin"]), space_after=0)
        if op.get("address"):
            label = "Address: " if op.get("is_individual") else "Registered Office: "
            para(doc, label + op["address"], space_after=0)
        contact = []
        if op.get("email"):
            contact.append("Email: {}".format(op["email"]))
        if op.get("phone"):
            contact.append("Phone: {}".format(op["phone"]))
        if contact:
            para(doc, " | ".join(contact), space_after=0)
        if not op.get("is_individual"):
            para(doc, "Through its Authorized Officers and Directors", space_after=0)
        para(doc, op_label(i, len(ops)), bold=True,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)


def bordered_table(doc, headers, rows):
    """A Table-Grid table. headers: list[str] or None. rows: list[list[str]]."""
    ncols = len(headers) if headers else len(rows[0])
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    if headers:
        hr = table.add_row().cells
        for c, h in enumerate(headers):
            hr[c].text = ""
            run = hr[c].paragraphs[0].add_run(h)
            run.bold = True
            hr[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for c, val in enumerate(row):
            cells[c].text = str(val)
    return table


# --------------------------------------------------------------------------
# cross-reference computation  (the anti-drift core)
# --------------------------------------------------------------------------
def complaint_para_count(case):
    """N = body paras + cause of action + limitation + jurisdiction paras + grounds para."""
    k = len(case["complaint_paragraphs"])
    j = len(case.get("jurisdiction_paragraphs", []))
    return k + 1 + 1 + j + 1  # +cause +limitation +jurisdiction(j) +grounds


def annexure_range(case):
    anx = case.get("annexures", [])
    if not anx:
        return None
    return "{} to {}".format(anx[0]["id"], anx[-1]["id"])


# --------------------------------------------------------------------------
# the six builders
# --------------------------------------------------------------------------
def build_index(doc, case):
    cause_title(doc, case)
    para(doc, "INDEX", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    pr = INDEX_PAGE_RANGES
    rows = [
        ["1", "Index", pr["index"]],
        ["2", "Proforma for Filing Consumer Complaint", pr["proforma"]],
        ["3", "Synopsis with List of Dates and Events", pr["synopsis"]],
        ["4", "Memo of Parties", pr["memo"]],
        ["5", "Consumer Complaint under Section 35 of the Consumer Protection Act, 2019 with Verification", pr["complaint"]],
        ["6", "Affidavit of the Complainant (notarised)", pr["affidavit"]],
    ]
    if case.get("annexures"):
        rows.append(["ANNEXURES", "", ""])
    for anx in case.get("annexures", []):
        rows.append([anx["id"], anx["description"], ""])
    bordered_table(doc, ["S.NO.", "PARTICULARS OF DOCUMENT", "PAGE NO."], rows)
    para(doc, "", space_after=2)
    signature_block(doc, case)


def build_proforma(doc, case):
    cause_title(doc, case)
    para(doc, "PROFORMA FOR FILING CONSUMER COMPLAINT", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    cm = case["complainant"]
    m = case["money"]
    ops = case["opposite_parties"]
    op1 = ops[0]
    pairs = [
        ("Name of the Complainant", cm["name"]),
        ("Father’s Name", cm["father_name"]),
        ("Age", "{} years".format(cm["age"])),
        ("Address", cm["address"]),
        ("Email & Mobile", "{} | {}".format(cm["email"], cm["mobile"])),
        ("Whether the Complainant is a Consumer within the meaning of Section 2(7) of the Consumer Protection Act, 2019",
         case["consumer_status"]["narrative"]),
        ("Name of Opposite Party No. 1", op1["name"] + ((" (" + op1["operating_as"] + ")") if op1.get("operating_as") else "")),
    ]
    if op1.get("cin"):
        pairs.append(("CIN of OP-1", op1["cin"]))
    pairs.append(("Address of OP-1", op1.get("address", "")))
    contact = []
    if op1.get("email"):
        contact.append(op1["email"])
    if op1.get("phone"):
        contact.append(op1["phone"])
    pairs.append(("Contact of OP-1", " | ".join(contact)))
    for extra in ops[1:]:
        pairs.append(("Name of Opposite Party No. {}".format(ops.index(extra) + 1),
                      extra["name"] + ((" — " + extra.get("address", "")) if extra.get("address") else "")))
    pairs += [
        ("Nature of Complaint", case["nature_of_complaint"]),
        ("Brief Facts", case["brief_facts"]),
        ("Cause of Action", case["cause_of_action"]),
        ("Limitation", case["limitation"]),
        ("Pecuniary Jurisdiction",
         "The value of the consideration paid (₹ {:,}/-) is below ₹ 50,00,000/-. "
         "This Hon’ble District Commission has pecuniary jurisdiction under Section 34(1) of the Act."
         .format(m["consideration_paid"])),
        ("Territorial Jurisdiction", case["territorial_jurisdiction"]),
        ("Amount paid as consideration",
         "₹ {:,}/- ({})".format(m["consideration_paid"], m["consideration_words"])),
        ("Total Claim Value (including compensation, interest and costs)",
         "₹ {:,}/- ({})".format(m["total_claim"], m["total_claim_words"])),
        ("Relief Sought", "; ".join("({}) {}".format(chr(97 + i), p.rstrip("; ."))
                                    for i, p in enumerate(case["prayer"])) + "."),
        ("Whether the matter has been referred to any other Court or Forum",
         "No. A grievance has been registered with the National Consumer Helpline "
         "(Docket No. {} dated {}), which is a mediation channel and not a judicial forum. "
         "No other Court or Commission is seized of the matter."
         .format(case["nch"]["docket"], case["nch"]["date"])),
        ("Whether the Complainant prays for ex-parte ad-interim relief",
         "No interim relief sought at this stage."),
    ]
    rows = [[str(i + 1), label, value] for i, (label, value) in enumerate(pairs)]
    bordered_table(doc, ["S.NO.", "PARTICULARS", "DETAILS"], rows)
    para(doc, "", space_after=2)
    signature_block(doc, case)


def build_synopsis(doc, case):
    cause_title(doc, case)
    para(doc, "SYNOPSIS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, case["synopsis"], align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)
    para(doc, "LIST OF DATES AND EVENTS", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    rows = [[d["date"], d["event"]] for d in case["dates_and_events"]]
    bordered_table(doc, ["DATE", "EVENT"], rows)
    para(doc, "", space_after=2)
    signature_block(doc, case)


def build_memo(doc, case):
    cause_title(doc, case)
    para(doc, "MEMO OF PARTIES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    party_block(doc, case)
    para(doc, "", space_after=2)
    signature_block(doc, case)


def build_complaint_affidavit(doc, case):
    cm = case["complainant"]
    N = complaint_para_count(case)
    anx_range = annexure_range(case)

    # --- the complaint ---
    cause_title(doc, case)
    party_block(doc, case)
    para(doc, "CONSUMER COMPLAINT UNDER SECTION 35 OF THE CONSUMER PROTECTION ACT, 2019",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, "RESPECTFULLY SHOWETH:", bold=True, space_after=6)

    n = 0
    for body in case["complaint_paragraphs"]:
        n += 1
        para(doc, "{}. {}".format(n, body), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    para(doc, "CAUSE OF ACTION:", bold=True, space_after=4)
    n += 1
    para(doc, "{}. {}".format(n, case["cause_of_action"]), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    para(doc, "LIMITATION:", bold=True, space_after=4)
    n += 1
    para(doc, "{}. {}".format(n, case["limitation"]), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    para(doc, "JURISDICTION:", bold=True, space_after=4)
    for jp in case.get("jurisdiction_paragraphs", []):
        n += 1
        para(doc, "{}. {}".format(n, jp), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    para(doc, "GROUNDS:", bold=True, space_after=4)
    n += 1
    para(doc, "{}. {}".format(n, case.get("grounds_intro", "That the conduct of the Opposite Parties constitutes:")),
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for gi, ground in enumerate(case["grounds"]):
        marker = ROMAN[gi] if gi < len(ROMAN) else str(gi + 1)
        para(doc, "({}) {}".format(marker, ground), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    assert n == N, "para-count drift: counted {} but computed {}".format(n, N)

    para(doc, "PRAYER:", bold=True, space_after=4)
    para(doc, "In the premises aforesaid, the Complainant most respectfully prays that this "
              "Hon’ble Commission may be pleased to:", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for pi, prayer in enumerate(case["prayer"]):
        para(doc, "({}) {}".format(chr(97 + pi), prayer), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    para(doc, "AND FOR THIS ACT OF KINDNESS THE COMPLAINANT, AS IN DUTY BOUND, SHALL EVER PRAY.",
         space_after=12)
    signature_block(doc, case)

    # --- verification ---
    para(doc, "VERIFICATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "I, {}, the Complainant above-named, do hereby solemnly verify that the contents "
              "of paragraphs 1 to {} of this Complaint are true and correct to my personal "
              "knowledge, and that the contents of the Prayer have been incorporated upon legal "
              "advice and belief. Nothing material has been concealed therefrom."
              .format(cm["name"], N), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    month = case["filing"].get("signing_month", "") or "__________"
    para(doc, "Verified at {} on this _____ day of {}, {}."
              .format(cm["city"], month, case["filing"]["year"]), space_after=12)
    para(doc, "({})".format(cm["name"]), align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(doc, "Complainant", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    # --- affidavit (new page) ---
    doc.add_page_break()
    cause_title(doc, case)
    para(doc, "AFFIDAVIT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, "I, {}, son of Mr. {}, aged about {} years, by faith {}, presently residing at "
              "{}, do hereby solemnly affirm and state on oath as follows:"
              .format(cm["name"], cm["father_name"], cm["age"], cm["faith"], cm["address"]),
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    note = case.get("affidavit_annexure_note", "") or ""
    true_copies = ("That the Annexures {} produced with the Complaint are true and faithful "
                   "copies of their respective originals{}.").format(
                       anx_range, (", " + note) if note else "")
    averments = [
        "That I am the Complainant in the accompanying Consumer Complaint and I am well "
        "acquainted with the facts and circumstances of the case and competent to swear this Affidavit.",
        "That I have read and understood the contents of the accompanying Consumer Complaint, "
        "the Memo of Parties, the Synopsis with List of Dates and Events, the Proforma for "
        "Filing Consumer Complaint, the Index, and the Annexures {} thereto.".format(anx_range),
        "That the statements made in paragraphs 1 to {} of the accompanying Consumer Complaint "
        "are true and correct to my personal knowledge.".format(N),
        "That the contents of the Synopsis and the List of Dates and Events accompanying this "
        "Complaint are true and correct to my personal knowledge.",
        true_copies,
        "That no part of this Affidavit is false and nothing material has been concealed therefrom.",
    ]
    for i, av in enumerate(averments, start=1):
        para(doc, "{}. {}".format(i, av), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    signature_block(doc, case, role="DEPONENT")

    para(doc, "VERIFICATION OF AFFIDAVIT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "Verified at {} on this _____ day of {}, {}, that the contents of the above "
              "Affidavit are true and correct to my personal knowledge, no part of it is false, "
              "and nothing material has been concealed therefrom."
              .format(cm["city"], month, case["filing"]["year"]), space_after=12)
    para(doc, "({})".format(cm["name"]), align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    para(doc, "DEPONENT", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    para(doc, "Identified by me:", space_after=18)
    para(doc, "Advocate / Notary Stamp & Seal Below:", space_after=0)


def build_party_in_person(doc, case):
    cm = case["complainant"]
    cause_title(doc, case)
    para(doc, "IN THE MATTER OF:", bold=True, space_after=8)
    para(doc, cm["name"], space_after=0)
    para(doc, "R/o {}".format(cm["address"]), space_after=0)
    para(doc, "… COMPLAINANT", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    para(doc, "VERSUS", bold=True, space_after=6)
    ops = case["opposite_parties"]
    for i, op in enumerate(ops, start=1):
        line = "{}. {}".format(i, op["name"])
        if op.get("operating_as"):
            line += " ({})".format(op["operating_as"])
        para(doc, line, space_after=0)
        if op.get("cin"):
            para(doc, "   CIN: {}".format(op["cin"]), space_after=0)
        if op.get("address"):
            para(doc, "   {}".format(op["address"]), space_after=4)
    para(doc, "… OPPOSITE {}".format("PARTY" if len(ops) == 1 else "PARTIES"),
         bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    para(doc, "DECLARATION OF PARTY-IN-PERSON", bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, "I, {}, the Complainant above-named, do hereby declare as follows:".format(cm["name"]),
         space_after=6)
    decls = [
        "That I am filing the accompanying Consumer Complaint, Memo of Parties, Synopsis, List "
        "of Dates and Events, Proforma and Affidavit, together with the documents listed in the "
        "Index, in person and without engaging an Advocate to represent me before this Hon’ble Commission.",
        "That I have decided to appear and act in person, and that I am fully aware of (a) the "
        "procedure of this Hon’ble Commission, (b) the consequences of so appearing in person, "
        "and (c) the requirement to make myself available for any hearing dates, additional "
        "documentation, or clarifications that this Hon’ble Commission may seek.",
        "That all communications, notices, orders and process from this Hon’ble Commission may "
        "be addressed to me at the address shown in the Memo of Parties or by email to {} and by "
        "mobile to {}.".format(cm["email"], cm["mobile"]),
        "That this declaration is made bona fide and not for any oblique purpose.",
    ]
    for i, d in enumerate(decls, start=1):
        para(doc, "{}. {}".format(i, d), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    signature_block(doc, case, role="Complainant, In Person")


# --------------------------------------------------------------------------
# orchestration
# --------------------------------------------------------------------------
REQUIRED_TOP_KEYS = ["output", "filing", "commission", "complainant", "opposite_parties",
                     "consumer_status", "money", "complaint_paragraphs", "cause_of_action",
                     "limitation", "grounds", "prayer", "annexures", "synopsis",
                     "dates_and_events", "nch"]


def validate(case):
    missing = [k for k in REQUIRED_TOP_KEYS if k not in case]
    if missing:
        sys.exit("Case YAML is missing required keys: {}".format(", ".join(missing)))
    if not case["opposite_parties"]:
        sys.exit("Case YAML has no opposite_parties.")
    # annexure ids should be sequential A-1, A-2, ...
    for i, anx in enumerate(case["annexures"], start=1):
        if anx.get("id") != "A-{}".format(i):
            sys.exit("Annexure ids must be sequential A-1, A-2, ...  found {!r} at position {}"
                     .format(anx.get("id"), i))


def main():
    ap = argparse.ArgumentParser(description="Generate a consumer-court filing pack from a case YAML.")
    ap.add_argument("case_yaml", help="path to the case spec YAML")
    ap.add_argument("--out", help="override output directory")
    args = ap.parse_args()

    with open(args.case_yaml, "r", encoding="utf-8") as fh:
        case = yaml.safe_load(fh)
    validate(case)

    out_dir = args.out or case["output"]["dir"]
    out_dir = os.path.abspath(out_dir)
    os.makedirs(out_dir, exist_ok=True)
    slug = case["output"]["case_slug"]
    start = int(case["output"].get("start_number", 1))

    specs = [
        ("Index", build_index),
        ("Proforma", build_proforma),
        ("Synopsis_and_Dates", build_synopsis),
        ("Memo_of_Parties", build_memo),
        ("Consumer_Complaint_with_Affidavit", build_complaint_affidavit),
    ]
    if case.get("party_in_person", True):
        specs.append(("Party_In_Person_Declaration", build_party_in_person))

    written = []
    for i, (label, builder) in enumerate(specs):
        doc = new_doc()
        builder(doc, case)
        fname = "{:02d}_{}_{}.docx".format(start + i, slug, label)
        path = os.path.join(out_dir, fname)
        doc.save(path)
        written.append(path)

    N = complaint_para_count(case)
    print("Generated {} files in {}".format(len(written), out_dir))
    for p in written:
        print("  " + os.path.basename(p))
    print()
    print("Cross-references derived by the generator:")
    print("  complaint paragraphs: 1 to {}".format(N))
    print("  annexures: {}".format(annexure_range(case)))
    print()
    print("NEXT: verify legal substance, confirm statute citations against")
    print("references/cp-act-2019.md, then notarise file 05 and file on e-Jagriti.")


if __name__ == "__main__":
    main()
